import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
import logging

# Configure logging
logger = logging.getLogger(__name__)

def parse_markdown_to_docx(text, paragraph):
    """Parses Markdown bold (**text**) into DOCX formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])  # Remove ** markers
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'

def parse_markdown_to_pdf(text):
    """Converts Markdown bold (**text**) into <b>text</b> for ReportLab."""
    return re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

def export_to_word(text, figure_data=None):
    """Export text and optional figure(s) to Word document."""
    doc = Document()
    
    # Add heading with custom styling
    heading = doc.add_heading("Research Summary", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(16)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue color
    heading.runs[0].font.name = 'Arial'
    
    # Add main text with Markdown bold formatting
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
    parse_markdown_to_docx(text, paragraph)
    
    # Process figures if provided
    if figure_data:
        try:
            figures = figure_data if isinstance(figure_data, list) else [figure_data]
            for fig in figures:
                buf = io.BytesIO()
                try:
                    # Handle matplotlib figure
                    if hasattr(fig, 'savefig'):
                        fig.savefig(buf, format='png')
                        buf.seek(0)
                        doc.add_picture(buf, width=Inches(5))
                    # Handle bytes or bytearray
                    elif isinstance(fig, (bytes, bytearray)):
                        buf = io.BytesIO(fig)
                        doc.add_picture(buf, width=Inches(5))
                    # Handle dictionary-wrapped figure
                    elif isinstance(fig, dict) and 'fig' in fig and hasattr(fig['fig'], 'savefig'):
                        fig['fig'].savefig(buf, format='png')
                        buf.seek(0)
                        doc.add_picture(buf, width=Inches(5))
                    else:
                        logger.warning(f"Skipped unsupported figure type: {type(fig)}")
                finally:
                    buf.close()
        except Exception as e:
            logger.error(f"Error processing figures for Word: {e}")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(text, figure_data=None):
    """Export text and optional figure(s) to PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=1*inch,
        bottomMargin=1*inch
    )
    
    # Define custom styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='CustomTitle',
        fontSize=16,
        leading=20,
        textColor=HexColor('#003366'),  # Dark blue
        fontName='Helvetica-Bold',
        alignment=1,  # Center
        spaceAfter=12
    ))
    styles.add(ParagraphStyle(
        name='CustomBody',
        fontSize=11,
        leading=14,
        fontName='Times-Roman',
        textColor=HexColor('#333333'),
        spaceAfter=10,
        justification=0  # Justified
    ))
    
    # Parse Markdown for PDF
    formatted_text = parse_markdown_to_pdf(text)
    
    # Build flowables
    flowables = [
        Paragraph("Research Summary", styles['CustomTitle']),
        Spacer(1, 0.2*inch),
        Paragraph(formatted_text, styles['CustomBody'])
    ]
    
    # Process figures if provided
    if figure_data:
        try:
            figures = figure_data if isinstance(figure_data, list) else [figure_data]
            for fig in figures:
                buf = io.BytesIO()
                try:
                    # Handle matplotlib figure
                    if hasattr(fig, 'savefig'):
                        fig.savefig(buf, format='png')
                        buf.seek(0)
                        flowables.append(Spacer(1, 0.3*inch))
                        flowables.append(RLImage(buf, width=5*inch, height=3*inch))
                    # Handle bytes or bytearray
                    elif isinstance(fig, (bytes, bytearray)):
                        buf = io.BytesIO(fig)
                        flowables.append(Spacer(1, 0.3*inch))
                        flowables.append(RLImage(buf, width=5*inch, height=3*inch))
                    # Handle dictionary-wrapped figure
                    elif isinstance(fig, dict) and 'fig' in fig and hasattr(fig['fig'], 'savefig'):
                        fig['fig'].savefig(buf, format='png')
                        buf.seek(0)
                        flowables.append(Spacer(1, 0.3*inch))
                        flowables.append(RLImage(buf, width=5*inch, height=3*inch))
                    else:
                        logger.warning(f"Skipped unsupported figure type: {type(fig)}")
                finally:
                    buf.close()
        except Exception as e:
            logger.error(f"Error processing figures for PDF: {e}")
    
    # Build PDF
    doc.build(flowables)
    buffer.seek(0)
    return buffer
